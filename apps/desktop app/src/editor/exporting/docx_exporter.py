"""DOCX exporter built on top of :mod:`python-docx`."""

from __future__ import annotations

from pathlib import Path

from .base import ExportError, ExportRequest, Exporter
from .markdown_blocks import BlockType, parse_markdown_to_blocks


class DocxExporter(Exporter):
    """Export documents as `.docx` files using :mod:`python-docx`."""

    def export(self, request: ExportRequest, target_path: Path) -> None:  # pragma: no cover - UI wiring
        text = (request.markdown or "").strip()
        if not text:
            raise ExportError("The document is empty; there is nothing to export.")

        try:
            import docx  # type: ignore[import]
        except Exception as exc:  # pragma: no cover - optional dependency
            raise ExportError("DOCX export is not available because 'python-docx' is not installed.") from exc

        try:
            document = docx.Document()

            # Map simple block structure to DOCX paragraphs/headings.
            for block in parse_markdown_to_blocks(request.markdown):
                if block.type is BlockType.HEADING:
                    level = block.level or 1
                    level = max(1, min(level, 4))
                    document.add_heading(block.text, level=level)
                else:
                    document.add_paragraph(block.text)

            # Basic metadata.
            props = document.core_properties
            if request.title:
                try:
                    props.title = request.title
                except Exception:
                    pass

            metadata = request.metadata or {}
            author = metadata.get("author") if isinstance(metadata, dict) else None
            if author:
                try:
                    props.author = str(author)
                except Exception:
                    pass

            document.save(str(target_path))
        except Exception as exc:
            raise ExportError(f"Failed to export DOCX: {exc}") from exc
